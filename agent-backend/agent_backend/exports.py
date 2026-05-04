"""Word-document exporters for BCG U Studio.

Three .docx endpoints today:
  - /export/script-docx       — Synthesia avatar script (Phase 1 #4j)
  - /export/case-study-docx   — case study handout (Phase 1 #5j)
  - /export/course-docx       — full course bundle (Phase 1 #6, in progress)

All three share:
  - Brand-driven section accents — the active brand drives every
    primary/dark/ink/inkLt color via _palette() (AI-1f). The brand
    toggle in the UI was previously surface-only; AI-1f wires it
    through to exports so a Client-branded course exports in blue,
    BCG in bright green, BCG U in deeper green — TOC, section
    headers, KC bullets, case study attribution, all of it.
  - Trebuchet MS body font, with the Henderson Sans upgrade path
    documented in _DOCX_FONT.
  - Consolas only on monospace columns (script SPOKEN cells).

This module exposes a FastAPI APIRouter; main.py wires it into the
app via include_router.
"""
from __future__ import annotations

import contextvars
import base64
import io
import logging
import re
import ssl

import certifi
import httpx

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Cm, Pt, RGBColor

_log = logging.getLogger(__name__)

# KK: shared SSL context for the banner fetch — same shape used in
# images.py's Pexels client. Falls back to certifi if truststore
# isn't available.
def _build_banner_ssl_context() -> ssl.SSLContext:
    try:
        import truststore  # type: ignore[import-not-found]
        return truststore.SSLContext(ssl.PROTOCOL_TLS_CLIENT)
    except ImportError:
        pass
    except Exception:
        pass
    try:
        return ssl.create_default_context(cafile=certifi.where())
    except Exception:
        return ssl.create_default_context()


_BANNER_SSL_CONTEXT = _build_banner_ssl_context()


router = APIRouter(prefix="/export", tags=["export"])


# ─── Brand-driven color palettes (AI-1f) ──────────────────────────────────────

# Per-brand palette dictionaries. Keys match BrandKey in
# app/src/brand/tokens.ts so server-side rendering matches what the
# in-app preview shows. Each export endpoint resolves the active
# brand into the _current_brand ContextVar at the top of the request,
# and adapters read colors via _palette().

_BRAND_PALETTES: dict[str, dict[str, RGBColor]] = {
    "bcg": {
        "primary":    RGBColor(0x29, 0xBA, 0x74),  # bright BCG green (B[bcg].pri)
        "primary_dk": RGBColor(0x1B, 0x7A, 0x4F),  # B[bcg].priDk
        "ink":        RGBColor(0x33, 0x33, 0x33),
        "ink_lt":     RGBColor(0x66, 0x66, 0x66),
    },
    "bcgu": {
        "primary":    RGBColor(0x19, 0x7A, 0x56),  # darker BCG U green (B[bcgu].pri)
        "primary_dk": RGBColor(0x0D, 0x3B, 0x2C),  # B[bcgu].priDk
        "ink":        RGBColor(0x33, 0x33, 0x33),
        "ink_lt":     RGBColor(0x66, 0x66, 0x66),
    },
    "client": {
        "primary":    RGBColor(0x25, 0x63, 0xEB),  # generic blue (B[client].pri)
        "primary_dk": RGBColor(0x1D, 0x4E, 0xD8),  # B[client].priDk
        "ink":        RGBColor(0x1E, 0x29, 0x3B),
        "ink_lt":     RGBColor(0x64, 0x74, 0x8B),
    },
}

# ContextVar threads the active brand through the render path. Each
# export endpoint calls `_current_brand.set(course.brand)` at the
# top; downstream adapters read via _palette() without needing to
# pass colors through every signature. Default = "bcgu" so tests +
# tools that import this module without setting the var still work.
_current_brand: contextvars.ContextVar[str] = contextvars.ContextVar(
    "_current_brand", default="bcgu",
)


def _palette() -> dict[str, RGBColor]:
    """Return the currently active brand's color palette.

    Falls back to bcgu if the active brand isn't recognized — handles
    legacy course payloads where brand might be missing or set to a
    deprecated value.
    """
    brand = _current_brand.get()
    return _BRAND_PALETTES.get(brand, _BRAND_PALETTES["bcgu"])


# Legacy color names — kept as functions returning the active palette
# so the 75 existing call sites in this module work unchanged. The
# BCG_ prefix is misleading post-AI-1f (these resolve per-brand) but
# renaming would churn every adapter; semantic comment here clarifies.
def _BCG_GREEN() -> RGBColor:  # noqa: N802 — legacy name, brand-resolved
    return _palette()["primary"]


def _BCG_INK() -> RGBColor:  # noqa: N802 — legacy name, brand-resolved
    return _palette()["ink"]


def _BCG_INK_LT() -> RGBColor:  # noqa: N802 — legacy name, brand-resolved
    return _palette()["ink_lt"]

# Default body font for every .docx export. Trebuchet MS is the
# Windows-built-in fallback BCG sanctions when the licensed Henderson
# Sans typeface isn't installed. Henderson upgrade path: drop the
# Henderson Sans .woff2 files into agent-backend/fonts/, swap this
# constant to "Henderson Sans", and add the install step to RUN.md.
# Out of scope for the pilot.
_DOCX_FONT = "Trebuchet MS"


def _set_docx_default_font(doc: Document) -> None:
    """Set the document-wide default font to _DOCX_FONT.

    Two layers needed:
      1. Normal style's rFonts — covers paragraphs styled as Normal
         (the default), which is most of our content.
      2. docDefaults rPrDefault rFonts — covers runs that don't
         inherit a style. Without this, Word's theme cascade falls
         back to Calibri (minorHAnsi theme) for stray runs, even
         though the style says Trebuchet.

    Runs that explicitly set run.font.name (e.g. "Consolas" for
    monospace columns) keep their override.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Layer 1 — Normal style.
    doc.styles["Normal"].font.name = _DOCX_FONT

    # Layer 2 — docDefaults / rPrDefault / rFonts.
    styles_el = doc.styles.element
    rpr_default = styles_el.find(qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr"))
    if rpr_default is None:
        return
    rfonts = rpr_default.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr_default.insert(0, rfonts)
    # Override theme-based attributes with explicit font names.
    for theme_attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        if rfonts.get(qn(f"w:{theme_attr}")) is not None:
            del rfonts.attrib[qn(f"w:{theme_attr}")]
    rfonts.set(qn("w:ascii"), _DOCX_FONT)
    rfonts.set(qn("w:hAnsi"), _DOCX_FONT)
    rfonts.set(qn("w:cs"), _DOCX_FONT)


def _safe_filename(stem: str) -> str:
    cleaned = re.sub(r"[^\w\-_.]", "_", stem).strip("_") or "export"
    return cleaned[:80]


# ─── Synthesia script exporter ────────────────────────────────────────────────


class ScriptExportRequest(BaseModel):
    script: str
    videoType: str = "speaker"
    lessonRef: str = ""
    courseName: str = ""
    duration: str = ""


def _parse_scenes(script: str) -> list[dict]:
    """Mirror of the FE parser. Returns a list of {index, spoken, visual}."""
    if not script.strip():
        return []
    if not re.search(r"SCENE\s+\d+", script, re.I):
        return []
    if not re.search(r"(SPOKEN|VISUAL):", script, re.I):
        return []

    scenes: list[dict] = []
    current: dict | None = None
    field: str | None = None
    for line in script.splitlines():
        m = re.match(r"^\s*SCENE\s+(\d+)", line, re.I)
        if m:
            if current:
                scenes.append(current)
            current = {"index": int(m.group(1)), "spoken": "", "visual": ""}
            field = None
            continue
        if not current:
            continue
        m = re.match(r"^\s*SPOKEN:\s*(.*)$", line, re.I)
        if m:
            field = "spoken"
            current["spoken"] = m.group(1)
            continue
        m = re.match(r"^\s*VISUAL:\s*(.*)$", line, re.I)
        if m:
            field = "visual"
            current["visual"] = m.group(1)
            continue
        if field and line.strip():
            sep = "\n" if current[field] else ""
            current[field] = current[field] + sep + line.strip()
    if current:
        scenes.append(current)
    return scenes


def _count_spoken_words(script: str) -> int:
    """Count words across all SPOKEN: blocks, stripping XML-ish tags."""
    matches = re.findall(
        r"SPOKEN:\s*([\s\S]*?)(?=\n\s*(?:VISUAL:|SCENE\s+\d+|$))",
        script,
        re.I,
    )
    spoken = " ".join(matches)
    cleaned = re.sub(r"<[^>]*>", "", spoken).strip()
    if not cleaned:
        return 0
    return len(re.split(r"\s+", cleaned))


@router.post("/script-docx")
async def export_script_docx(req: ScriptExportRequest):
    if not req.script or not req.script.strip():
        raise HTTPException(status_code=400, detail="script is required")

    scenes = _parse_scenes(req.script)
    word_count = _count_spoken_words(req.script)
    seconds = round(word_count / 150 * 60) if word_count else 0

    doc = Document()
    _set_docx_default_font(doc)

    # Title — course name
    title = doc.add_paragraph()
    run = title.add_run(req.courseName or "Synthesia Script")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _BCG_GREEN()

    # Subtitle — lesson / type / duration
    sub_bits = []
    if req.lessonRef:
        sub_bits.append(f"Lesson {req.lessonRef}")
    if req.videoType:
        sub_bits.append(req.videoType.capitalize())
    if req.duration:
        sub_bits.append(req.duration)
    if sub_bits:
        sub = doc.add_paragraph()
        run = sub.add_run(" · ".join(sub_bits))
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_INK_LT()

    doc.add_paragraph()  # spacer

    if scenes:
        # Section header for the table
        header = doc.add_paragraph()
        run = header.add_run("Scenes")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_GREEN()

        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        # Header row
        hdr_cells = table.rows[0].cells
        for i, label in enumerate(["#", "Spoken", "Visual"]):
            cell = hdr_cells[i]
            p = cell.paragraphs[0]
            r = p.add_run(label)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_GREEN()

        # Approximate column widths.
        widths_cm = (1.0, 9.0, 6.0)
        for col_idx, w in enumerate(widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)

        for s in scenes:
            row = table.add_row().cells
            # # (scene number)
            p = row[0].paragraphs[0]
            r = p.add_run(str(s["index"]))
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK()
            # Spoken — monospace for readability
            p = row[1].paragraphs[0]
            r = p.add_run(s["spoken"])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK()
            # Visual
            p = row[2].paragraphs[0]
            r = p.add_run(s["visual"])
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK_LT()
    else:
        # Fallback — script didn't parse as scenes; embed raw.
        header = doc.add_paragraph()
        run = header.add_run("Script (raw — couldn't parse as scenes)")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_GREEN()

        body = doc.add_paragraph()
        run = body.add_run(req.script)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    # Footer — word count + duration estimate
    doc.add_paragraph()
    foot = doc.add_paragraph()
    run = foot.add_run(f"~{word_count} words · ~{seconds} sec at 150 wpm")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _BCG_INK_LT()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    stem = (
        _safe_filename(f"{req.courseName}-{req.lessonRef}-script")
        if req.courseName
        else _safe_filename(f"{req.lessonRef}-script")
    )
    filename = f"{stem}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Case study exporter ──────────────────────────────────────────────────────


class CaseStudyStakeholderModel(BaseModel):
    name: str = ""
    role: str = ""
    voice: str = ""


class CaseStudyModel(BaseModel):
    id: str = ""
    title: str = ""
    context: str = ""
    stakeholders: list[CaseStudyStakeholderModel] = []
    decisionPoints: list[str] = []
    debriefPrompts: list[str] = []


class CaseStudyExportRequest(BaseModel):
    caseStudy: CaseStudyModel
    courseName: str = ""
    moduleTitle: str = ""


def _split_sources(context: str) -> tuple[str, str | None]:
    """Pull out a trailing 'Sources' / 'Inspired by' section if present.

    The Case Study Designer prompt asks the agent to append a brief
    Sources block at the end of context when materials are attached.
    We render that as its own styled section in the .docx instead of
    leaving it inline at the bottom of the context paragraph.
    """
    pattern = re.compile(
        r"\n\s*(?:#+\s*)?(?:Sources|Inspired by)[:\s]*\n",
        re.I,
    )
    m = pattern.search(context)
    if not m:
        return context, None
    body = context[: m.start()].rstrip()
    sources = context[m.end():].strip()
    return body, sources or None


@router.post("/case-study-docx")
async def export_case_study_docx(req: CaseStudyExportRequest):
    cs = req.caseStudy
    if not cs.context.strip() and not cs.stakeholders:
        raise HTTPException(status_code=400, detail="case study has no content to export")

    body_context, sources_block = _split_sources(cs.context)

    doc = Document()
    _set_docx_default_font(doc)

    # --- Title block ---
    title = doc.add_paragraph()
    run = title.add_run(req.courseName or "Case Study")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _BCG_GREEN()

    sub_bits = []
    if req.moduleTitle:
        sub_bits.append(req.moduleTitle)
    sub_bits.append(cs.title or "Untitled case study")
    sub = doc.add_paragraph()
    run = sub.add_run(" · ".join(sub_bits))
    run.font.size = Pt(11)
    run.font.color.rgb = _BCG_INK_LT()

    doc.add_paragraph()  # spacer

    def section_heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_GREEN()

    def body_paragraph(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_INK()

    # --- Context ---
    if body_context.strip():
        section_heading("Context")
        for para in [p for p in body_context.split("\n\n") if p.strip()]:
            body_paragraph(para.strip())

    # --- Stakeholders ---
    if cs.stakeholders:
        section_heading("Stakeholders")
        for s in cs.stakeholders:
            p = doc.add_paragraph()
            r = p.add_run(s.name)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK()
            if s.role:
                r2 = p.add_run(f" — {s.role}")
                r2.font.size = Pt(11)
                r2.font.color.rgb = _BCG_INK_LT()
            if s.voice:
                quote = doc.add_paragraph()
                rq = quote.add_run(f"“{s.voice}”")
                rq.italic = True
                rq.font.size = Pt(10)
                rq.font.color.rgb = _BCG_INK()
                quote.paragraph_format.left_indent = Cm(0.6)

    # --- Decision points ---
    if cs.decisionPoints:
        section_heading("Decision points")
        for i, dp in enumerate(cs.decisionPoints, start=1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. ")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_GREEN()
            r2 = p.add_run(dp)
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK()

    # --- Debrief prompts ---
    if cs.debriefPrompts:
        section_heading("Debrief prompts (for facilitation)")
        for i, dp in enumerate(cs.debriefPrompts, start=1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. ")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_GREEN()
            r2 = p.add_run(dp)
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK()

    # --- Sources ---
    if sources_block:
        section_heading("Sources / Inspired by")
        for line in [l for l in sources_block.split("\n") if l.strip()]:
            cleaned = re.sub(r"^[\-\*•]\s*", "", line.strip())
            p = doc.add_paragraph()
            r = p.add_run(f"• {cleaned}")
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK_LT()

    # --- Stream out ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    stem = _safe_filename(
        f"{req.courseName}-{req.moduleTitle}-case-study"
        if req.courseName or req.moduleTitle
        else f"{cs.title}-case-study"
    )
    filename = f"{stem}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Course .docx exporter (Phase 1 #6) ───────────────────────────────────────

# Block-type adapters for Word doc rendering
# ──────────────────────────────────────────
# text         → paragraph
# banner       → styled section header + body paragraph
# callout      → bordered tinted block (label + body)
# cards        → bulleted list of card titles + descriptions
# accordion    → expanded headings + body (collapsible-equivalent)
# flipcard     → "Front: X / Back: Y" pairs
# timeline     → numbered ordered list (each item: title + description)
# quiz         → question + options with correct marker + rationale
# poll         → question + options with % values
# stats        → number + label list (KPI-style)
# image        → caption + alt text only (no embed for v1)
# video        → embedded script in 3-col table (reuses script renderer)
# divider      → SKIPPED (genuinely contentless)
# ──────────────────────────────────────────
# If you add a new block type and don't add an adapter here, the export
# will fall back to a generic placeholder line. Don't let that happen
# — add the row.


class BlockItemModel(BaseModel):
    title: str = ""
    desc: str | None = None
    img: str | None = None
    alt: str | None = None


class BlockDataModel(BaseModel):
    content: str | None = None
    url: str | None = None
    caption: str | None = None
    alt: str | None = None
    title: str | None = None
    body: str | None = None
    type: str | None = None
    items: list[BlockItemModel] | None = None
    script: str | None = None
    videoType: str | None = None
    # AI-1b additions — must match BlockData in app/src/course/types.ts
    # so server-side .docx export sees the same fields the FE writes.
    imageUrl: str | None = None              # banner statement-mode photo
    attribution: str | None = None           # quote speaker name
    attributionRole: str | None = None       # quote role/company
    attributionPhotoUrl: str | None = None   # quote round-photo URL
    iconName: str | None = None              # sectionHeader icon name


class BlockModel(BaseModel):
    id: str
    type: str
    data: BlockDataModel = BlockDataModel()


class QuizQuestionModel(BaseModel):
    type: str  # "mcq" | "short"
    stem: str = ""
    options: list[str] | None = None
    correctIndex: int | None = None
    rationale: str | None = None
    expectedAnswerHints: list[str] | None = None


class QuizModel(BaseModel):
    questions: list[QuizQuestionModel] = []


class LessonModel(BaseModel):
    id: str
    title: str = ""
    duration: int = 10
    blocks: list[BlockModel] = []
    objectives: list[str] | None = None
    knowledgeCheck: QuizModel | None = None
    # KK: lesson hero banner. Embedded at the top of each lesson's
    # .docx section as a full-width image with photographer credit
    # underneath. Mirrors the FE Lesson type fields; safe to omit.
    bannerImageUrl: str | None = None
    bannerPhotographer: str | None = None
    bannerPhotographerUrl: str | None = None


class CourseModuleModel(BaseModel):
    id: str
    title: str = ""
    weekNumber: int | None = None
    summary: str | None = None
    objectives: list[str] | None = None
    knowledgeCheck: QuizModel | None = None
    caseStudyId: str | None = None
    lessons: list[LessonModel] = []


class MaterialModel(BaseModel):
    id: str = ""
    filename: str = ""
    charCount: int = 0


class CourseModel(BaseModel):
    id: str
    title: str = ""
    client: str = ""
    brand: str = "bcgu"
    modules: list[CourseModuleModel] = []
    materials: list[MaterialModel] | None = None
    caseStudies: list[CaseStudyModel] | None = None


class CourseExportRequest(BaseModel):
    course: CourseModel
    audience: str = ""


# ─── Section helpers ──────────────────────────────────────────────────────────


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = _BCG_GREEN()


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = _BCG_GREEN()


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = _BCG_GREEN()


def _body(doc: Document, text: str, italic: bool = False, light: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.italic = italic
    r.font.color.rgb = _BCG_INK_LT() if light else _BCG_INK()


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"• {text}")
    r.font.size = Pt(11)
    r.font.color.rgb = _BCG_INK()


def _numbered(doc: Document, idx: int, text: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{idx}. ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = _BCG_GREEN()
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = _BCG_INK()


def _page_break(doc: Document) -> None:
    doc.add_page_break()


# ─── Block-type adapters ──────────────────────────────────────────────────────


def _render_text_block(doc: Document, block: BlockModel) -> None:
    """Render a text block to .docx with markdown bold support (AI-1a).

    Lesson Writer v2 (AI-1c) emits body paragraphs with strategic inline
    `**bold**` markers — concepts, action verbs, named entities, decision
    phrasing. We split each paragraph on the bold markers and emit
    separate Runs with bold=True for the wrapped segments. Plain runs
    inherit Trebuchet MS body (set on the Normal style).
    """
    content = (block.data.content or "").strip()
    if not content:
        return
    for para in [p.strip() for p in content.split("\n\n") if p.strip()]:
        _body_with_bold(doc, para)


# Markdown bold pattern — matches the JS-side renderInlineMd in
# app/src/course/previewBlock.ts. Non-greedy, no internal asterisks.
_BOLD_RE = re.compile(r"\*\*([^*]+?)\*\*")


def _body_with_bold(doc: Document, text: str) -> None:
    """Add a body paragraph, splitting on **markdown bold** markers.

    Plain segments and bold segments emit separate runs in the same
    paragraph. Run-level bold doesn't change font size/family — those
    inherit from the Normal style (Trebuchet MS, set in
    _set_docx_default_font).
    """
    p = doc.add_paragraph()
    last_end = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last_end:
            r = p.add_run(text[last_end:m.start()])
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK()
        r = p.add_run(m.group(1))
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = _BCG_INK()
        last_end = m.end()
    if last_end < len(text):
        r = p.add_run(text[last_end:])
        r.font.size = Pt(11)
        r.font.color.rgb = _BCG_INK()


def _render_banner_block(doc: Document, block: BlockModel) -> None:
    """Render a banner block to .docx (AI-1b extended).

    The banner's optional imageUrl (statement-mode background photo)
    isn't embedded in the .docx — fetching remote URLs at export time
    would balloon doc size + hit network reliability questions in
    enterprise environments. The .docx version surfaces title + body
    only; the photo lives in the web preview / SCORM export.
    """
    title = (block.data.title or "").strip()
    body = (block.data.body or "").strip()
    if title:
        _h3(doc, title)
    if body:
        # AI-1f: banner body now respects **markdown bold** runs so the
        # statement's emphasized phrases land bolded in print.
        _body_with_bold(doc, body)


def _render_callout_block(doc: Document, block: BlockModel) -> None:
    """Render a callout block to .docx with markdown bold body (AI-1f).

    Callout variants (info / tip / note / warning / success) prefix the
    body with `[VARIANT]` in primary color, then the body text in ink
    with **markdown bold** runs split via _body_with_bold.
    """
    callout_kind = (block.data.type or "tip").upper()
    body = (block.data.body or "").strip()
    if not body:
        return
    # Title row with [VARIANT] prefix in primary color.
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{callout_kind}] ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = _BCG_GREEN()
    # Body — split on **bold** markers, append runs to the same paragraph.
    last_end = 0
    for m in _BOLD_RE.finditer(body):
        if m.start() > last_end:
            r = p.add_run(body[last_end:m.start()])
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK()
        rb = p.add_run(m.group(1))
        rb.bold = True
        rb.font.size = Pt(11)
        rb.font.color.rgb = _BCG_INK()
        last_end = m.end()
    if last_end < len(body):
        r = p.add_run(body[last_end:])
        r.font.size = Pt(11)
        r.font.color.rgb = _BCG_INK()


# ─── AI-1b new block adapters ─────────────────────────────────────────────────


def _render_quote_block(doc: Document, block: BlockModel) -> None:
    """Render a quote block — italic body + attribution row (AI-1f)."""
    body = (block.data.body or "").strip()
    attribution = (block.data.attribution or "").strip()
    role = (block.data.attributionRole or "").strip()
    if not body:
        return
    # Quote body — italic, ink-900, with bold runs preserved.
    p = doc.add_paragraph()
    qopen = p.add_run('"')
    qopen.italic = True
    qopen.font.size = Pt(11)
    qopen.font.color.rgb = _BCG_INK()
    last_end = 0
    for m in _BOLD_RE.finditer(body):
        if m.start() > last_end:
            r = p.add_run(body[last_end:m.start()])
            r.italic = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK()
        rb = p.add_run(m.group(1))
        rb.italic = True
        rb.bold = True
        rb.font.size = Pt(11)
        rb.font.color.rgb = _BCG_INK()
        last_end = m.end()
    if last_end < len(body):
        r = p.add_run(body[last_end:])
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = _BCG_INK()
    qclose = p.add_run('"')
    qclose.italic = True
    qclose.font.size = Pt(11)
    qclose.font.color.rgb = _BCG_INK()
    # Attribution row — em-dash + name (bold, primary) + role (ink-lt).
    if attribution or role:
        a = doc.add_paragraph()
        a.paragraph_format.left_indent = Cm(0.6)
        em = a.add_run("— ")
        em.font.size = Pt(10)
        em.font.color.rgb = _BCG_INK_LT()
        if attribution:
            r = a.add_run(attribution)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_GREEN()
        if role:
            r = a.add_run(f", {role}" if attribution else role)
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK_LT()


def _render_click_instruction_block(doc: Document, block: BlockModel) -> None:
    """Render a clickInstruction block — short italic green hint (AI-1f).

    In .docx the click affordance doesn't apply (paper / Word doesn't
    have an interactive equivalent), but we still emit the hint as an
    italic green callout so SME reviewers reading the doc see "this is
    where the learner gets a click cue" — useful for review even
    though it doesn't drive interaction.
    """
    content = (block.data.content or "").strip()
    if not content:
        return
    p = doc.add_paragraph()
    r1 = p.add_run("→ ")
    r1.italic = True
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = _BCG_GREEN()
    r2 = p.add_run(content)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = _BCG_GREEN()


def _render_section_header_block(doc: Document, block: BlockModel) -> None:
    """Render a sectionHeader block — icon glyph + title divider (AI-1f).

    .docx doesn't render lucide icons natively; we use a unicode glyph
    fallback that matches the previewBlock.ts iconGlyphs map (so web
    preview + .docx export agree on the visual). The title gets h3
    styling with the brand-primary color, prefixed by the icon glyph.
    """
    title = (block.data.title or "").strip()
    icon_name = (block.data.iconName or "bookOpen").strip()
    if not title:
        return
    # Curated 12 icon glyphs — must match SECTION_ICON_NAMES + the
    # iconGlyphs map in app/src/course/previewBlock.ts. Out-of-set
    # names fall back to bookOpen.
    icon_glyphs: dict[str, str] = {
        "target":      "◎",
        "brain":       "🧠",
        "pencil":      "✎",
        "quote":       "❝",
        "check":       "✓",
        "clock":       "◷",
        "lightbulb":   "💡",
        "bookOpen":    "📖",
        "sparkles":    "✦",
        "alertCircle": "!",
        "trendingUp":  "↗",
        "users":       "👥",
    }
    glyph = icon_glyphs.get(icon_name, icon_glyphs["bookOpen"])
    p = doc.add_paragraph()
    r1 = p.add_run(f"{glyph}  ")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = _BCG_GREEN()
    r2 = p.add_run(title.upper())
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = _BCG_GREEN()


def _render_cards_block(doc: Document, block: BlockModel) -> None:
    items = block.data.items or []
    for item in items:
        title = (item.title or "").strip()
        desc = (item.desc or "").strip()
        if not title and not desc:
            continue
        p = doc.add_paragraph()
        if title:
            r = p.add_run(f"• {title}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK()
        if desc:
            r2 = p.add_run(f"  — {desc}" if title else f"• {desc}")
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK_LT()


def _render_accordion_block(doc: Document, block: BlockModel) -> None:
    items = block.data.items or []
    for item in items:
        title = (item.title or "").strip()
        desc = (item.desc or "").strip()
        if title:
            _h3(doc, title)
        if desc:
            _body(doc, desc)


def _render_flipcard_block(doc: Document, block: BlockModel) -> None:
    items = block.data.items or []
    for item in items:
        front = (item.title or "").strip()
        back = (item.desc or "").strip()
        if not front and not back:
            continue
        p = doc.add_paragraph()
        r1 = p.add_run("Front: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = _BCG_GREEN()
        r2 = p.add_run(front)
        r2.font.size = Pt(11)
        r2.font.color.rgb = _BCG_INK()
        if back:
            p2 = doc.add_paragraph()
            r3 = p2.add_run("Back: ")
            r3.bold = True
            r3.font.size = Pt(11)
            r3.font.color.rgb = _BCG_GREEN()
            r4 = p2.add_run(back)
            r4.font.size = Pt(11)
            r4.font.color.rgb = _BCG_INK()


def _render_timeline_block(doc: Document, block: BlockModel) -> None:
    items = block.data.items or []
    for i, item in enumerate(items, start=1):
        title = (item.title or "").strip()
        desc = (item.desc or "").strip()
        text = title if not desc else f"{title} — {desc}" if title else desc
        if text:
            _numbered(doc, i, text)


def _render_quiz_block(doc: Document, block: BlockModel) -> None:
    """Inline single-question quiz block (lesson body, not knowledge check)."""
    items = block.data.items or []
    if not items:
        return
    question = (items[0].title or "").strip()
    if question:
        _h3(doc, f"Q: {question}")
    for opt in items[1:]:
        title = (opt.title or "").strip()
        is_correct = (opt.desc or "") == "1"
        if not title:
            continue
        prefix = "✓ " if is_correct else "○ "
        p = doc.add_paragraph()
        r = p.add_run(prefix + title)
        r.font.size = Pt(11)
        r.font.color.rgb = _BCG_GREEN() if is_correct else _BCG_INK()
        r.bold = is_correct


def _render_poll_block(doc: Document, block: BlockModel) -> None:
    items = block.data.items or []
    if not items:
        return
    question = (items[0].title or "").strip()
    if question:
        _h3(doc, f"Poll: {question}")
    for opt in items[1:]:
        title = (opt.title or "").strip()
        pct = (opt.desc or "").strip()
        if not title:
            continue
        suffix = f" — {pct}%" if pct else ""
        _bullet(doc, f"{title}{suffix}")


def _render_stats_block(doc: Document, block: BlockModel) -> None:
    items = block.data.items or []
    for item in items:
        number = (item.title or "").strip()
        label = (item.desc or "").strip()
        if not number and not label:
            continue
        p = doc.add_paragraph()
        if number:
            r = p.add_run(number)
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = _BCG_GREEN()
        if label:
            r2 = p.add_run(f"  {label}")
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK_LT()


def _render_image_block(doc: Document, block: BlockModel) -> None:
    """v1: caption + alt only. Embedding the actual image is Phase 2."""
    caption = (block.data.caption or "").strip()
    alt = (block.data.alt or "").strip()
    url = (block.data.url or "").strip()
    p = doc.add_paragraph()
    r1 = p.add_run("[Image] ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = _BCG_GREEN()
    r2 = p.add_run(caption or alt or url or "(no caption / alt / url)")
    r2.font.size = Pt(11)
    r2.font.color.rgb = _BCG_INK_LT()
    r2.italic = True


def _render_video_block(doc: Document, block: BlockModel) -> None:
    """Reuses the script-table layout from /export/script-docx."""
    caption = (block.data.caption or "").strip()
    url = (block.data.url or "").strip()
    script = (block.data.script or "").strip()

    if caption or url:
        p = doc.add_paragraph()
        r1 = p.add_run("[Video] ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = _BCG_GREEN()
        r2 = p.add_run(caption or url)
        r2.font.size = Pt(11)
        r2.font.color.rgb = _BCG_INK_LT()
        r2.italic = True

    if not script:
        _body(doc, "(No Synthesia script attached.)", italic=True, light=True)
        return

    scenes = _parse_scenes(script)
    if not scenes:
        # Raw script fallback
        _body(doc, "Script (raw):", italic=True, light=True)
        body = doc.add_paragraph()
        r = body.add_run(script)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(["#", "Spoken", "Visual"]):
        p = hdr[i].paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = _BCG_GREEN()

    widths_cm = (1.0, 9.0, 6.0)
    for col_idx, w in enumerate(widths_cm):
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(w)

    for s in scenes:
        row = table.add_row().cells
        p = row[0].paragraphs[0]
        r = p.add_run(str(s["index"]))
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = _BCG_INK()
        p = row[1].paragraphs[0]
        r = p.add_run(s["spoken"])
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = _BCG_INK()
        p = row[2].paragraphs[0]
        r = p.add_run(s["visual"])
        r.font.size = Pt(10)
        r.font.color.rgb = _BCG_INK_LT()


_BLOCK_ADAPTERS: dict[str, callable] = {
    "text": _render_text_block,
    "banner": _render_banner_block,
    "callout": _render_callout_block,
    "cards": _render_cards_block,
    "accordion": _render_accordion_block,
    "flipcard": _render_flipcard_block,
    "timeline": _render_timeline_block,
    "quiz": _render_quiz_block,
    "poll": _render_poll_block,
    "stats": _render_stats_block,
    "image": _render_image_block,
    "video": _render_video_block,
    # AI-1b new block adapters (AI-1f wires them in).
    "quote": _render_quote_block,
    "clickInstruction": _render_click_instruction_block,
    "sectionHeader": _render_section_header_block,
    # divider: intentionally absent — skipped as genuinely contentless.
}


def _render_block(doc: Document, block: BlockModel) -> None:
    if block.type == "divider":
        return  # skip
    adapter = _BLOCK_ADAPTERS.get(block.type)
    if adapter is None:
        # Unknown block type — placeholder so the LD knows something
        # exists in the source course but didn't render. Add an
        # adapter to _BLOCK_ADAPTERS to remove the placeholder.
        p = doc.add_paragraph()
        r = p.add_run(f"[Block: {block.type} — content not yet rendered for print]")
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = _BCG_INK_LT()
        return
    adapter(doc, block)


# ─── Section sub-functions ────────────────────────────────────────────────────


def _render_cover(doc: Document, course: CourseModel, audience: str) -> None:
    # Eyebrow
    p = doc.add_paragraph()
    r = p.add_run("BCG U · COURSE")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _BCG_GREEN()

    _h1(doc, course.title or "Untitled Course")

    if audience:
        _body(doc, f"For: {audience}", italic=True, light=True)

    # Module / lesson totals
    module_count = len(course.modules)
    lesson_count = sum(len(m.lessons) for m in course.modules)
    _body(
        doc,
        f"{module_count} module{'s' if module_count != 1 else ''} · "
        f"{lesson_count} lesson{'s' if lesson_count != 1 else ''}",
        italic=True,
        light=True,
    )

    # Learning outcomes — accumulate every module's objectives.
    all_objectives: list[str] = []
    for m in course.modules:
        if m.objectives:
            all_objectives.extend(m.objectives)
    if all_objectives:
        doc.add_paragraph()  # spacer
        _h2(doc, "Learning outcomes")
        for o in all_objectives:
            _bullet(doc, o)


def _render_toc(doc: Document, course: CourseModel) -> None:
    _h2(doc, "Contents")
    for mi, m in enumerate(course.modules, start=1):
        # Module row
        p = doc.add_paragraph()
        r1 = p.add_run(f"Module {mi}. ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = _BCG_GREEN()
        r2 = p.add_run(m.title or "Untitled module")
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = _BCG_INK()
        # Lesson rows
        for li, lesson in enumerate(m.lessons, start=1):
            p2 = doc.add_paragraph()
            r1 = p2.add_run(f"        {mi}.{li}  ")
            r1.font.size = Pt(11)
            r1.font.color.rgb = _BCG_INK_LT()
            r2 = p2.add_run(lesson.title or "Untitled lesson")
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK()


def _render_knowledge_check(doc: Document, quiz: QuizModel, scope_label: str) -> None:
    _h3(doc, f"{scope_label} — {len(quiz.questions)} question{'s' if len(quiz.questions) != 1 else ''}")
    for i, q in enumerate(quiz.questions, start=1):
        # Track-R / R2: bump pre-question spacing so adjacent questions
        # don't visually run together. User feedback after the live KC
        # download: questions felt cramped. Adding ~14pt before each
        # stem (and an empty paragraph after the previous question's
        # rationale / hints) gives the page rhythm.
        if i > 1:
            spacer = doc.add_paragraph()
            spacer_run = spacer.add_run("")
            spacer_run.font.size = Pt(6)
        # Question stem
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(14)
        p_format.space_after = Pt(4)
        r1 = p.add_run(f"{i}. ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = _BCG_GREEN()
        r2 = p.add_run(q.stem)
        r2.font.size = Pt(11)
        r2.font.color.rgb = _BCG_INK()
        # Type tag
        type_label = "MCQ" if q.type == "mcq" else "Short answer"
        r3 = p.add_run(f"   ({type_label})")
        r3.italic = True
        r3.font.size = Pt(9)
        r3.font.color.rgb = _BCG_INK_LT()

        if q.type == "mcq":
            for oi, option in enumerate(q.options or []):
                is_correct = oi == (q.correctIndex if q.correctIndex is not None else -1)
                marker = "✓ " if is_correct else "○ "
                p2 = doc.add_paragraph()
                r = p2.add_run(f"   {marker}{option}")
                r.font.size = Pt(11)
                r.font.color.rgb = _BCG_GREEN() if is_correct else _BCG_INK()
                r.bold = is_correct
            if q.rationale:
                p3 = doc.add_paragraph()
                r1 = p3.add_run("Rationale: ")
                r1.bold = True
                r1.italic = True
                r1.font.size = Pt(10)
                r1.font.color.rgb = _BCG_INK_LT()
                r2 = p3.add_run(q.rationale)
                r2.italic = True
                r2.font.size = Pt(10)
                r2.font.color.rgb = _BCG_INK_LT()
        else:
            # Short answer — show expected hints (rubric).
            if q.expectedAnswerHints:
                p2 = doc.add_paragraph()
                r = p2.add_run("Expected answer hints (rubric):")
                r.bold = True
                r.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = _BCG_INK_LT()
                for hint in q.expectedAnswerHints:
                    _bullet(doc, hint)


def _fetch_banner_bytes(url: str) -> bytes | None:
    """KK: fetch the banner image bytes for embedding into the .docx.

    Handles both data: URLs (LD-uploaded files, where the image is
    inline-encoded) and remote HTTPS URLs (Pexels / Unsplash). Returns
    None on any failure — caller falls through to "no banner"
    rendering so the export never breaks because of a missing image.

    OO3: every silent-failure path now logs a warning so future
    regressions surface in the backend log instead of producing a
    silently-image-less docx.
    """
    if not url:
        return None
    if url.startswith("data:"):
        try:
            header, _, b64 = url.partition(",")
            if not b64 or "base64" not in header:
                _log.warning("KK banner: malformed data URL (header=%s)", header[:50])
                return None
            return base64.b64decode(b64)
        except Exception as exc:
            _log.warning("KK banner: data-URL decode failed: %s", exc)
            return None
    try:
        with httpx.Client(timeout=8.0, verify=_BANNER_SSL_CONTEXT) as client:
            resp = client.get(url, follow_redirects=True)
            resp.raise_for_status()
            return resp.content
    except Exception as exc:
        _log.warning("KK banner: fetch %s failed: %s", url[:80], exc)
        return None


def _pexels_lookup_first(query: str) -> tuple[str, str, str] | None:
    """OO3: server-side fallback Pexels search for lessons that don't
    have a bannerImageUrl persisted yet.

    The auto-fetcher in LessonBanner.tsx only fires when the lesson
    canvas mounts. If an LD generates a course but never opens a
    given lesson, that lesson's banner is never persisted — so the
    .docx export sees `bannerImageUrl=None` and skips the embed.

    To make the export self-sufficient, we run the same Pexels
    search on the server when bannerImageUrl is missing. Returns
    (url, photographer, photographer_url) on success, or None when:
      - PEXELS_API_KEY isn't configured (graceful no-op)
      - the query is empty
      - Pexels returns 0 results / errors

    Logs at warning level on every failure path so the docx export
    surfaces a missing-banner regression in the backend log.
    """
    from agent_backend.images import (
        PEXELS_API_KEY,
        PEXELS_BASE,
        _PEXELS_SSL_CONTEXT,
    )

    if not query.strip():
        return None
    if not PEXELS_API_KEY:
        _log.warning(
            "KK banner: lesson has no bannerImageUrl AND PEXELS_API_KEY "
            "is unset — docx will render without a banner. Set the key "
            "in agent-backend/.env to enable server-side fallback."
        )
        return None
    try:
        with httpx.Client(timeout=10.0, verify=_PEXELS_SSL_CONTEXT) as client:
            resp = client.get(
                f"{PEXELS_BASE}/search",
                params={"query": query, "per_page": 1, "orientation": "landscape"},
                headers={"Authorization": PEXELS_API_KEY},
            )
            resp.raise_for_status()
            data = resp.json()
    except Exception as exc:
        _log.warning("KK banner fallback: Pexels lookup failed for %r: %s", query, exc)
        return None
    photos = data.get("photos") or []
    if not photos:
        _log.warning("KK banner fallback: Pexels returned 0 results for %r", query)
        return None
    p = photos[0]
    url = (p.get("src") or {}).get("large") or ""
    if not url:
        _log.warning("KK banner fallback: Pexels first result has no large src")
        return None
    return (
        url,
        p.get("photographer") or "",
        p.get("photographer_url") or "",
    )


def _render_lesson_banner(doc: Document, lesson: LessonModel) -> None:
    """KK + OO3: embed the lesson's banner photo at the top of its
    .docx section, sized to the page width, with a small attribution
    caption underneath when the photographer is known.

    OO3: when the lesson has no bannerImageUrl persisted (e.g. the LD
    never opened the lesson canvas, so the auto-fetcher never fired),
    fall back to a synchronous Pexels lookup using the lesson title
    as the query. This makes the .docx export self-sufficient — it
    doesn't depend on the LD having opened every lesson.

    Silent on failure (with log.warning at every drop point) so the
    lesson body still renders cleanly even when banners can't load.
    """
    url = lesson.bannerImageUrl
    photographer = lesson.bannerPhotographer or ""
    photographer_url = lesson.bannerPhotographerUrl or ""

    if not url:
        # OO3: fall back to a server-side Pexels lookup so lessons the
        # LD never opened still ship with a banner.
        fallback = _pexels_lookup_first(lesson.title or "")
        if fallback:
            url, photographer, photographer_url = fallback
            _log.info(
                "KK banner fallback: server-fetched banner for %r",
                lesson.title[:60] if lesson.title else "",
            )
        else:
            return  # nothing to embed; logged inside _pexels_lookup_first

    img_bytes = _fetch_banner_bytes(url)
    if not img_bytes:
        return
    try:
        # Page text width on A4 with 2.5cm side margins is ~16cm.
        doc.add_picture(io.BytesIO(img_bytes), width=Cm(16))
    except Exception as exc:
        _log.warning("KK banner: add_picture failed (%d bytes): %s", len(img_bytes), exc)
        return
    # Suppress the photographer_url variable lint when only the
    # photographer name is used in the caption — the URL is reserved
    # for a future hyperlink upgrade. Reference it explicitly so
    # linters don't strip the binding.
    _ = photographer_url
    if photographer:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"Photo by {photographer} on Pexels")
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = _BCG_INK_LT()


def _render_lesson(doc: Document, lesson: LessonModel, mi: int, li: int) -> None:
    _render_lesson_banner(doc, lesson)
    _h2(doc, f"{mi}.{li}  {lesson.title or 'Untitled lesson'}")
    _body(doc, f"{lesson.duration} min · {len(lesson.blocks)} block{'s' if len(lesson.blocks) != 1 else ''}", italic=True, light=True)

    if lesson.objectives:
        _h3(doc, "Lesson objectives")
        for o in lesson.objectives:
            _bullet(doc, o)
        doc.add_paragraph()  # spacer

    # Blocks in order, via per-type adapters.
    for block in lesson.blocks:
        _render_block(doc, block)

    # Lesson knowledge check (if any).
    if lesson.knowledgeCheck and lesson.knowledgeCheck.questions:
        doc.add_paragraph()
        _render_knowledge_check(doc, lesson.knowledgeCheck, "Knowledge check")


def _render_module(doc: Document, module: CourseModuleModel, mi: int, course: CourseModel) -> None:
    _page_break(doc)
    week_label = f"Week {module.weekNumber}" if module.weekNumber else f"Module {mi}"
    p = doc.add_paragraph()
    r = p.add_run(f"{week_label.upper()} · MODULE")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _BCG_GREEN()

    _h1(doc, module.title or "Untitled module")

    if module.summary:
        _body(doc, module.summary, italic=True, light=True)

    if module.objectives:
        doc.add_paragraph()
        _h3(doc, "Module objectives")
        for o in module.objectives:
            _bullet(doc, o)

    # Lessons.
    for li, lesson in enumerate(module.lessons, start=1):
        doc.add_paragraph()
        _render_lesson(doc, lesson, mi, li)

    # Module final assessment.
    if module.knowledgeCheck and module.knowledgeCheck.questions:
        doc.add_paragraph()
        _render_knowledge_check(doc, module.knowledgeCheck, "Final assessment")

    # Case study reference.
    if module.caseStudyId and course.caseStudies:
        cs = next((c for c in course.caseStudies if c.id == module.caseStudyId), None)
        if cs:
            doc.add_paragraph()
            _h3(doc, f"Case study: {cs.title}")
            _body(doc, "(See case study section below.)", italic=True, light=True)


def _render_case_study(doc: Document, cs: CaseStudyModel, course_title: str) -> None:
    _page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("CASE STUDY")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _BCG_GREEN()

    _h1(doc, cs.title or "Untitled case study")

    body_context, sources_block = _split_sources(cs.context)

    if body_context.strip():
        _h3(doc, "Context")
        for para in [p for p in body_context.split("\n\n") if p.strip()]:
            _body(doc, para.strip())

    if cs.stakeholders:
        _h3(doc, "Stakeholders")
        for s in cs.stakeholders:
            p = doc.add_paragraph()
            r = p.add_run(s.name)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK()
            if s.role:
                r2 = p.add_run(f" — {s.role}")
                r2.font.size = Pt(11)
                r2.font.color.rgb = _BCG_INK_LT()
            if s.voice:
                quote = doc.add_paragraph()
                rq = quote.add_run(f"“{s.voice}”")
                rq.italic = True
                rq.font.size = Pt(10)
                rq.font.color.rgb = _BCG_INK()
                quote.paragraph_format.left_indent = Cm(0.6)

    if cs.decisionPoints:
        _h3(doc, "Decision points")
        for i, dp in enumerate(cs.decisionPoints, start=1):
            _numbered(doc, i, dp)

    if cs.debriefPrompts:
        _h3(doc, "Debrief prompts (for facilitation)")
        for i, dp in enumerate(cs.debriefPrompts, start=1):
            _numbered(doc, i, dp)

    if sources_block:
        _h3(doc, "Sources / Inspired by")
        for line in [l for l in sources_block.split("\n") if l.strip()]:
            cleaned = re.sub(r"^[\-\*•]\s*", "", line.strip())
            _body(doc, f"• {cleaned}", light=True)


def _render_appendix(doc: Document, course: CourseModel) -> None:
    materials = course.materials or []
    if not materials:
        return
    _page_break(doc)
    _h1(doc, "Source materials")
    _body(
        doc,
        "Files the LD ingested while authoring this course. Lesson "
        "writers and the case study designer may have drawn on these "
        "as source material.",
        italic=True,
        light=True,
    )
    doc.add_paragraph()
    for m in materials:
        p = doc.add_paragraph()
        r1 = p.add_run("• ")
        r1.font.size = Pt(11)
        r1.font.color.rgb = _BCG_GREEN()
        r2 = p.add_run(m.filename or "(unnamed file)")
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = _BCG_INK()
        if m.charCount:
            r3 = p.add_run(f"   ({m.charCount:,} chars)")
            r3.font.size = Pt(10)
            r3.font.color.rgb = _BCG_INK_LT()


# ─── Endpoint ─────────────────────────────────────────────────────────────────


@router.post("/course-docx")
async def export_course_docx(req: CourseExportRequest):
    course = req.course
    if not course.modules:
        raise HTTPException(status_code=400, detail="course has no modules to export")

    # AI-1f: thread the course's brand into the render path. ContextVar
    # is request-scoped (FastAPI gives each request its own task; tasks
    # have isolated ContextVar reads). All adapters downstream call
    # _palette() which reads this ContextVar, so brand colors flow
    # through TOC + section headers + KC bullets + case study
    # attribution + every block adapter without arg-threading.
    _current_brand.set(course.brand or "bcgu")

    doc = Document()
    _set_docx_default_font(doc)

    _render_cover(doc, course, req.audience)
    _page_break(doc)
    _render_toc(doc, course)

    for mi, m in enumerate(course.modules, start=1):
        _render_module(doc, m, mi, course)

    # Case studies as their own end-of-doc section. Each gets its own page.
    for cs in course.caseStudies or []:
        # Skip empty slots (planted but never designed).
        if not cs.context.strip() and not cs.stakeholders:
            continue
        _render_case_study(doc, cs, course.title)

    _render_appendix(doc, course)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    stem = _safe_filename(f"{course.title or 'course'}-course")
    filename = f"{stem}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Track-B: standalone KC export ────────────────────────────────────────────


class KcExportRequest(BaseModel):
    """Track-B (KC Studio): standalone knowledge-check download.

    The Kc record is FE-side localStorage-persisted; FE POSTs the
    full record to this endpoint when the LD clicks Download. The
    rendering reuses _render_knowledge_check (same path as the
    course-docx export's lesson + module KC sections), wrapped in
    a focused single-page document with a brand-tinted header.
    """
    title: str = ""
    topic: str = ""
    questions: list[QuizQuestionModel] = []
    brand: str | None = None


@router.post("/kc-docx")
async def export_kc_docx(req: KcExportRequest):
    if not req.questions:
        raise HTTPException(status_code=400, detail="kc has no questions to export")

    _current_brand.set(req.brand or "bcgu")

    doc = Document()
    _set_docx_default_font(doc)

    # Header — title + topic + question count.
    title_text = req.title or req.topic or "Knowledge check"
    p = doc.add_paragraph()
    r = p.add_run("KNOWLEDGE CHECK")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _BCG_GREEN()
    _h1(doc, title_text)
    if req.topic and req.topic != title_text:
        _body(doc, f"Topic: {req.topic}", italic=True, light=True)
    _body(
        doc,
        f"{len(req.questions)} question{'s' if len(req.questions) != 1 else ''}",
        italic=True,
        light=True,
    )
    doc.add_paragraph()

    # Reuse the canonical KC renderer with a generic scope label.
    _render_knowledge_check(doc, QuizModel(questions=req.questions), "Questions")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    stem = _safe_filename(f"{title_text}-kc")
    filename = f"{stem}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
