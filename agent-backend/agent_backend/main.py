import io
import logging
import re

from fastapi import FastAPI, HTTPException, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Cm, Pt, RGBColor

from . import config  # noqa: F401 — loads .env and sets git-bash env var on import
from .parse import ParseError, SUPPORTED_EXTENSIONS, parse_file
from .session import Session

MAX_UPLOAD_BYTES = 25 * 1024 * 1024  # 25 MB hard cap per file

logging.basicConfig(level=logging.INFO)
log = logging.getLogger(__name__)

app = FastAPI(title="NovoEd Course Builder Agent Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[config.ALLOWED_ORIGIN],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health():
    return {"ok": True}


@app.post("/parse")
async def parse_endpoint(file: UploadFile):
    filename = file.filename or "upload"
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"file too large: {len(data) // 1024} KB exceeds {MAX_UPLOAD_BYTES // (1024 * 1024)} MB cap",
        )
    try:
        text = parse_file(filename, data)
    except ParseError as exc:
        raise HTTPException(status_code=415, detail=str(exc))
    return {
        "filename": filename,
        "text": text,
        "charCount": len(text),
        "supported": SUPPORTED_EXTENSIONS,
    }


class ScriptExportRequest(BaseModel):
    script: str
    videoType: str = "speaker"
    lessonRef: str = ""
    courseName: str = ""
    duration: str = ""


# BCG green for section headers / accents in the exported .docx.
_BCG_GREEN = RGBColor(0x1B, 0x7A, 0x4F)
_BCG_INK = RGBColor(0x33, 0x33, 0x33)
_BCG_INK_LT = RGBColor(0x66, 0x66, 0x66)

# Default body font for every .docx export. Trebuchet MS is the
# Windows-built-in fallback BCG sanctions when the licensed Henderson
# Sans typeface isn't installed. Henderson upgrade path: drop the
# Henderson Sans .woff2 files into agent-backend/fonts/, swap this
# constant to "Henderson Sans", and add the install step to RUN.md.
# Out of scope for the pilot.
_DOCX_FONT = "Trebuchet MS"


def _set_docx_default_font(doc: Document) -> None:
    """Set the document-wide default font to _DOCX_FONT.

    Two layers needed:
      1. Normal style's rFonts — covers paragraphs styled as Normal
         (the default), which is most of our content.
      2. docDefaults rPrDefault rFonts — covers runs that don't
         inherit a style. Without this, Word's theme cascade falls
         back to Calibri (minorHAnsi theme) for stray runs, even
         though the style says Trebuchet.

    Runs that explicitly set run.font.name (e.g. "Consolas" for
    monospace columns) keep their override.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Layer 1 — Normal style.
    doc.styles["Normal"].font.name = _DOCX_FONT

    # Layer 2 — docDefaults / rPrDefault / rFonts.
    styles_el = doc.styles.element
    rpr_default = styles_el.find(qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr"))
    if rpr_default is None:
        return
    rfonts = rpr_default.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr_default.insert(0, rfonts)
    # Override theme-based attributes with explicit font names.
    for theme_attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        if rfonts.get(qn(f"w:{theme_attr}")) is not None:
            del rfonts.attrib[qn(f"w:{theme_attr}")]
    rfonts.set(qn("w:ascii"), _DOCX_FONT)
    rfonts.set(qn("w:hAnsi"), _DOCX_FONT)
    rfonts.set(qn("w:cs"), _DOCX_FONT)


def _parse_scenes(script: str) -> list[dict]:
    """Mirror of the FE parser. Returns a list of {index, spoken, visual}."""
    if not script.strip():
        return []
    if not re.search(r"SCENE\s+\d+", script, re.I):
        return []
    if not re.search(r"(SPOKEN|VISUAL):", script, re.I):
        return []

    scenes: list[dict] = []
    current: dict | None = None
    field: str | None = None
    for line in script.splitlines():
        m = re.match(r"^\s*SCENE\s+(\d+)", line, re.I)
        if m:
            if current:
                scenes.append(current)
            current = {"index": int(m.group(1)), "spoken": "", "visual": ""}
            field = None
            continue
        if not current:
            continue
        m = re.match(r"^\s*SPOKEN:\s*(.*)$", line, re.I)
        if m:
            field = "spoken"
            current["spoken"] = m.group(1)
            continue
        m = re.match(r"^\s*VISUAL:\s*(.*)$", line, re.I)
        if m:
            field = "visual"
            current["visual"] = m.group(1)
            continue
        if field and line.strip():
            sep = "\n" if current[field] else ""
            current[field] = current[field] + sep + line.strip()
    if current:
        scenes.append(current)
    return scenes


def _count_spoken_words(script: str) -> int:
    """Count words across all SPOKEN: blocks, stripping XML-ish tags."""
    matches = re.findall(
        r"SPOKEN:\s*([\s\S]*?)(?=\n\s*(?:VISUAL:|SCENE\s+\d+|$))",
        script,
        re.I,
    )
    spoken = " ".join(matches)
    cleaned = re.sub(r"<[^>]*>", "", spoken).strip()
    if not cleaned:
        return 0
    return len(re.split(r"\s+", cleaned))


def _safe_filename(stem: str) -> str:
    cleaned = re.sub(r"[^\w\-_.]", "_", stem).strip("_") or "script"
    return cleaned[:80]


@app.post("/export/script-docx")
async def export_script_docx(req: ScriptExportRequest):
    if not req.script or not req.script.strip():
        raise HTTPException(status_code=400, detail="script is required")

    scenes = _parse_scenes(req.script)
    word_count = _count_spoken_words(req.script)
    seconds = round(word_count / 150 * 60) if word_count else 0

    doc = Document()
    _set_docx_default_font(doc)

    # Title — course name
    title = doc.add_paragraph()
    run = title.add_run(req.courseName or "Synthesia Script")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _BCG_GREEN

    # Subtitle — lesson / type / duration
    sub_bits = []
    if req.lessonRef:
        sub_bits.append(f"Lesson {req.lessonRef}")
    if req.videoType:
        sub_bits.append(req.videoType.capitalize())
    if req.duration:
        sub_bits.append(req.duration)
    if sub_bits:
        sub = doc.add_paragraph()
        run = sub.add_run(" · ".join(sub_bits))
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_INK_LT

    doc.add_paragraph()  # spacer

    if scenes:
        # Section header for the table
        header = doc.add_paragraph()
        run = header.add_run("Scenes")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_GREEN

        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        # Header row
        hdr_cells = table.rows[0].cells
        for i, label in enumerate(["#", "Spoken", "Visual"]):
            cell = hdr_cells[i]
            p = cell.paragraphs[0]
            r = p.add_run(label)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_GREEN

        # Approximate column widths.
        widths_cm = (1.0, 9.0, 6.0)
        for col_idx, w in enumerate(widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)

        for s in scenes:
            row = table.add_row().cells
            # # (scene number)
            p = row[0].paragraphs[0]
            r = p.add_run(str(s["index"]))
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK
            # Spoken — monospace for readability
            p = row[1].paragraphs[0]
            r = p.add_run(s["spoken"])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK
            # Visual
            p = row[2].paragraphs[0]
            r = p.add_run(s["visual"])
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK_LT
    else:
        # Fallback — script didn't parse as scenes; embed raw.
        header = doc.add_paragraph()
        run = header.add_run("Script (raw — couldn't parse as scenes)")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_GREEN

        body = doc.add_paragraph()
        run = body.add_run(req.script)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    # Footer — word count + duration estimate
    doc.add_paragraph()
    foot = doc.add_paragraph()
    run = foot.add_run(f"~{word_count} words · ~{seconds} sec at 150 wpm")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _BCG_INK_LT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    stem = _safe_filename(f"{req.courseName}-{req.lessonRef}-script") if req.courseName else _safe_filename(f"{req.lessonRef}-script")
    filename = f"{stem}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


class CaseStudyStakeholderModel(BaseModel):
    name: str = ""
    role: str = ""
    voice: str = ""


class CaseStudyModel(BaseModel):
    id: str = ""
    title: str = ""
    context: str = ""
    stakeholders: list[CaseStudyStakeholderModel] = []
    decisionPoints: list[str] = []
    debriefPrompts: list[str] = []


class CaseStudyExportRequest(BaseModel):
    caseStudy: CaseStudyModel
    courseName: str = ""
    moduleTitle: str = ""


def _split_sources(context: str) -> tuple[str, str | None]:
    """Pull out a trailing 'Sources' / 'Inspired by' section if present.

    The Case Study Designer prompt asks the agent to append a brief
    Sources block at the end of context when materials are attached.
    We render that as its own styled section in the .docx instead of
    leaving it inline at the bottom of the context paragraph.
    """
    pattern = re.compile(
        r"\n\s*(?:#+\s*)?(?:Sources|Inspired by)[:\s]*\n",
        re.I,
    )
    m = pattern.search(context)
    if not m:
        return context, None
    body = context[: m.start()].rstrip()
    sources = context[m.end():].strip()
    return body, sources or None


@app.post("/export/case-study-docx")
async def export_case_study_docx(req: CaseStudyExportRequest):
    cs = req.caseStudy
    if not cs.context.strip() and not cs.stakeholders:
        raise HTTPException(status_code=400, detail="case study has no content to export")

    body_context, sources_block = _split_sources(cs.context)

    doc = Document()
    _set_docx_default_font(doc)

    # --- Title block ---
    title = doc.add_paragraph()
    run = title.add_run(req.courseName or "Case Study")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _BCG_GREEN

    sub_bits = []
    if req.moduleTitle:
        sub_bits.append(req.moduleTitle)
    sub_bits.append(cs.title or "Untitled case study")
    sub = doc.add_paragraph()
    run = sub.add_run(" · ".join(sub_bits))
    run.font.size = Pt(11)
    run.font.color.rgb = _BCG_INK_LT

    doc.add_paragraph()  # spacer

    def section_heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_GREEN

    def body_paragraph(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = _BCG_INK

    # --- Context ---
    if body_context.strip():
        section_heading("Context")
        for para in [p for p in body_context.split("\n\n") if p.strip()]:
            body_paragraph(para.strip())

    # --- Stakeholders ---
    if cs.stakeholders:
        section_heading("Stakeholders")
        for s in cs.stakeholders:
            p = doc.add_paragraph()
            r = p.add_run(s.name)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_INK
            if s.role:
                r2 = p.add_run(f" — {s.role}")
                r2.font.size = Pt(11)
                r2.font.color.rgb = _BCG_INK_LT
            if s.voice:
                quote = doc.add_paragraph()
                rq = quote.add_run(f"“{s.voice}”")
                rq.italic = True
                rq.font.size = Pt(10)
                rq.font.color.rgb = _BCG_INK
                quote.paragraph_format.left_indent = Cm(0.6)

    # --- Decision points ---
    if cs.decisionPoints:
        section_heading("Decision points")
        for i, dp in enumerate(cs.decisionPoints, start=1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. ")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_GREEN
            r2 = p.add_run(dp)
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK

    # --- Debrief prompts ---
    if cs.debriefPrompts:
        section_heading("Debrief prompts (for facilitation)")
        for i, dp in enumerate(cs.debriefPrompts, start=1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. ")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _BCG_GREEN
            r2 = p.add_run(dp)
            r2.font.size = Pt(11)
            r2.font.color.rgb = _BCG_INK

    # --- Sources ---
    if sources_block:
        section_heading("Sources / Inspired by")
        for line in [l for l in sources_block.split("\n") if l.strip()]:
            cleaned = re.sub(r"^[\-\*•]\s*", "", line.strip())
            p = doc.add_paragraph()
            r = p.add_run(f"• {cleaned}")
            r.font.size = Pt(10)
            r.font.color.rgb = _BCG_INK_LT

    # --- Stream out ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    stem = _safe_filename(
        f"{req.courseName}-{req.moduleTitle}-case-study"
        if req.courseName or req.moduleTitle
        else f"{cs.title}-case-study"
    )
    filename = f"{stem}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.websocket("/ws")
async def ws_endpoint(websocket: WebSocket):
    await websocket.accept()
    session = Session(websocket)
    try:
        await session.start()
        while True:
            raw = await websocket.receive_text()
            await session.handle_client_message(raw)
    except WebSocketDisconnect:
        log.info("client disconnected")
    except Exception:
        log.exception("ws session error")
    finally:
        await session.close()
